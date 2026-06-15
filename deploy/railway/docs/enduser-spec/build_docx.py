#!/usr/bin/env python3
"""Assemble the NuFi end-user functional specification section files into a
single professional .docx. Reusable for English and Vietnamese builds."""

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BRAND = RGBColor(0x5B, 0x3D, 0xF5)      # NuFi purple-ish
DARK = RGBColor(0x1A, 0x1A, 0x2E)
GREY = RGBColor(0x60, 0x60, 0x70)
CODE_BG = "F2F2F7"
RULE_CLR = "D8D8E0"

# ---------------------------------------------------------------- low-level

def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_bottom_border(paragraph, color=RULE_CLR, sz="6"):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), sz)
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color)
    pbdr.append(bottom)
    pPr.append(pbdr)


def add_page_number_field(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)


INLINE_RE = re.compile(r'(\*\*.+?\*\*|`[^`]+?`)')

def add_inline(paragraph, text, base_size=None, base_color=None):
    """Render a markdown inline string with **bold** and `code` support."""
    text = text.replace(' ', ' ')
    for tok in INLINE_RE.split(text):
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**'):
            r = paragraph.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith('`') and tok.endswith('`'):
            r = paragraph.add_run(tok[1:-1]); r.font.name = 'Consolas'; r.font.size = Pt(9.5)
            rpr = r._r.get_or_add_rPr()
            shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), CODE_BG)
            rpr.append(shd)
        else:
            r = paragraph.add_run(tok)
        if base_size: r.font.size = base_size
        if base_color: r.font.color.rgb = base_color


# ---------------------------------------------------------------- styles

def setup_styles(doc):
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    heads = [('Heading 1', 17, BRAND, True, 14, 6),
             ('Heading 2', 13.5, DARK, True, 12, 4),
             ('Heading 3', 11.5, DARK, True, 9, 3),
             ('Heading 4', 10.5, GREY, True, 7, 2)]
    for name, size, color, bold, before, after in heads:
        st = doc.styles[name]
        st.font.size = Pt(size); st.font.color.rgb = color; st.font.bold = bold
        st.font.name = 'Calibri'
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True


# ---------------------------------------------------------------- markdown block parser

def render_table(doc, rows):
    # rows: list of list-of-cell-strings; first row header, drop separator row
    header = rows[0]
    body = [r for r in rows[1:] if not re.match(r'^\s*:?-{2,}', r[0].strip())]
    ncol = len(header)
    table = doc.add_table(rows=1, cols=ncol)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = table.rows[0].cells
    for i, c in enumerate(header):
        set_cell_bg(hdr[i], "EDE9FE")
        p = hdr[i].paragraphs[0]; p.paragraph_format.space_after = Pt(1)
        add_inline(p, c.strip())
        for run in p.runs: run.bold = True; run.font.size = Pt(9.5)
    for r in body:
        cells = table.add_row().cells
        for i in range(ncol):
            val = r[i].strip() if i < len(r) else ''
            p = cells[i].paragraphs[0]; p.paragraph_format.space_after = Pt(1)
            add_inline(p, val)
            for run in p.runs: run.font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def split_row(line):
    s = line.strip()
    if s.startswith('|'): s = s[1:]
    if s.endswith('|'): s = s[:-1]
    return [c for c in re.split(r'(?<!\\)\|', s)]


def render_markdown(doc, md_lines):
    i, n = 0, len(md_lines)
    while i < n:
        line = md_lines[i].rstrip('\n')
        stripped = line.strip()

        if not stripped:
            i += 1; continue

        # code fence
        if stripped.startswith('```'):
            i += 1; buf = []
            while i < n and not md_lines[i].strip().startswith('```'):
                buf.append(md_lines[i].rstrip('\n')); i += 1
            i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(6)
            r = p.add_run('\n'.join(buf)); r.font.name = 'Consolas'; r.font.size = Pt(9)
            rpr = r._r.get_or_add_rPr()
            shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), CODE_BG)
            rpr.append(shd)
            continue

        # table
        if stripped.startswith('|') and i + 1 < n and re.match(r'^\s*\|?\s*:?-{2,}', md_lines[i+1]):
            rows = []
            while i < n and md_lines[i].strip().startswith('|'):
                rows.append(split_row(md_lines[i])); i += 1
            render_table(doc, rows)
            continue

        # headings
        m = re.match(r'^(#{2,6})\s+(.*)', stripped)
        if m:
            level = len(m.group(1)) - 1   # ## -> H1
            level = min(level, 4)
            h = doc.add_heading(level=level)
            add_inline(h, m.group(2).strip())
            i += 1; continue

        # horizontal rule
        if re.match(r'^(\*\s*){3,}$', stripped) or re.match(r'^(-\s*){3,}$', stripped) or stripped in ('---', '***', '___'):
            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(4)
            add_bottom_border(p)
            i += 1; continue

        # blockquote
        if stripped.startswith('>'):
            buf = []
            while i < n and md_lines[i].strip().startswith('>'):
                buf.append(re.sub(r'^\s*>\s?', '', md_lines[i].rstrip('\n'))); i += 1
            text = ' '.join(x.strip() for x in buf if x.strip())
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(6)
            add_bottom_border(p, color="C7B8F5", sz="4")
            pPr = p._p.get_or_add_pPr()
            pbdr = pPr.find(qn('w:pBdr'))
            left = OxmlElement('w:left'); left.set(qn('w:val'),'single'); left.set(qn('w:sz'),'18'); left.set(qn('w:space'),'8'); left.set(qn('w:color'),'8B5CF6')
            pbdr.insert(0, left)
            add_inline(p, text, base_size=Pt(10), base_color=GREY)
            continue

        # bullet list
        mb = re.match(r'^(\s*)[-*+]\s+(.*)', line)
        if mb:
            indent = len(mb.group(1))
            lvl = min(indent // 2, 2)
            style = 'List Bullet' if lvl == 0 else f'List Bullet {lvl+1}'
            try:
                p = doc.add_paragraph(style=style)
            except KeyError:
                p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(2)
            add_inline(p, mb.group(2).strip())
            i += 1; continue

        # numbered list (keep literal number from source for stable numbering)
        mn = re.match(r'^(\s*)(\d+)\.\s+(.*)', line)
        if mn:
            indent = len(mn.group(1))
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25 + 0.25 * (indent // 2))
            p.paragraph_format.first_line_indent = Inches(-0.25)
            p.paragraph_format.space_after = Pt(2)
            num = p.add_run(f"{mn.group(2)}.  "); num.bold = True
            add_inline(p, mn.group(3).strip())
            i += 1; continue

        # plain paragraph
        p = doc.add_paragraph()
        add_inline(p, stripped)
        i += 1


# ---------------------------------------------------------------- doc assembly

def add_title_page(doc, S):
    for _ in range(4): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('NuFi'); r.font.size = Pt(40); r.font.bold = True; r.font.color.rgb = BRAND
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(S['title']); r.font.size = Pt(22); r.font.bold = True; r.font.color.rgb = DARK
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(S['subtitle']); r.font.size = Pt(13); r.font.color.rgb = GREY
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; add_bottom_border(p, color=str(hex(BRAND)[2:]).upper() if False else "8B5CF6", sz="8")
    for label, value in S['meta']:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rl = p.add_run(label + '  '); rl.bold = True; rl.font.size = Pt(10.5); rl.font.color.rgb = DARK
        rv = p.add_run(value); rv.font.size = Pt(10.5); rv.font.color.rgb = GREY
    doc.add_page_break()


def add_toc(doc, S):
    h = doc.add_heading(level=1); add_inline(h, S['toc_title'])
    p = doc.add_paragraph(); add_inline(p, S['toc_note'], base_size=Pt(9.5), base_color=GREY)
    para = doc.add_paragraph()
    run = para.add_run()
    fldBegin = OxmlElement('w:fldChar'); fldBegin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldSep = OxmlElement('w:fldChar'); fldSep.set(qn('w:fldCharType'), 'separate')
    t = OxmlElement('w:t'); t.text = S['toc_placeholder']
    fldEnd = OxmlElement('w:fldChar'); fldEnd.set(qn('w:fldCharType'), 'end')
    run._r.append(fldBegin); run._r.append(instr); run._r.append(fldSep); run._r.append(t); run._r.append(fldEnd)
    doc.add_page_break()


def add_part_divider(doc, label, title):
    doc.add_page_break()
    for _ in range(6): doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(label); r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = GREY
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title); r.font.size = Pt(26); r.font.bold = True; r.font.color.rgb = BRAND
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; add_bottom_border(p, color="8B5CF6", sz="10")


def setup_header_footer(doc, S):
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]; hp.text = ''
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = hp.add_run(S['running_header']); r.font.size = Pt(8); r.font.color.rgb = GREY
    footer = section.footer
    fp = footer.paragraphs[0]; fp.text = ''
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run(S['footer'] + '    |    '); r.font.size = Pt(8); r.font.color.rgb = GREY
    add_page_number_field(fp)
    for run in fp.runs: run.font.size = Pt(8); run.font.color.rgb = GREY


def build(sections_dir, output_path, S, manifest):
    doc = Document()
    setup_styles(doc)
    for s in doc.sections:
        s.top_margin = Inches(0.9); s.bottom_margin = Inches(0.9)
        s.left_margin = Inches(1.0); s.right_margin = Inches(1.0)
    add_title_page(doc, S)
    add_toc(doc, S)
    for entry in manifest:
        if entry[0] == 'part':
            add_part_divider(doc, entry[1], entry[2])
        else:
            path = Path(sections_dir) / entry[1]
            lines = path.read_text(encoding='utf-8').splitlines()
            render_markdown(doc, lines)
    setup_header_footer(doc, S)
    doc.save(output_path)
    print('Saved', output_path)


# ---------------------------------------------------------------- manifests

MANIFEST = [
    ('file', '00-introduction.md'),
    ('file', '00b-overview.md'),
    ('file', '00c-glossary.md'),
    ('part', 'PART A', 'NuFi Chat — End-User Features'),
    ('file', '01-auth.md'),
    ('file', '02-chat-core.md'),
    ('file', '03-endpoints-models-presets.md'),
    ('file', '04-agents-rag.md'),
    ('file', '05-files.md'),
    ('file', '06-conversation-mgmt.md'),
    ('file', '07-prompts.md'),
    ('file', '08-settings-console-link.md'),
    ('part', 'PART B', 'NuFi Console'),
    ('file', '09-console.md'),
    ('part', 'PART C', 'Cross-Cutting & Appendices'),
    ('file', '10-crosscutting.md'),
    ('file', '11-tester-guide.md'),
]

STRINGS_EN = {
    'title': 'NuFi Chat & Console',
    'subtitle': 'End-User Functional Specification (QA Reference)',
    'meta': [('Product:', 'NuFi Chat (LibreChat fork) + NuFi Console'),
             ('Audience:', 'QA / Testing'),
             ('Status:', 'Living document — v1.0 draft'),
             ('Date:', 'June 2026')],
    'toc_title': 'Table of Contents',
    'toc_note': 'This table of contents is a Word field. To populate page numbers, open in Microsoft Word / LibreOffice, select all (Ctrl+A) and update fields (F9), or right-click → "Update Field".',
    'toc_placeholder': 'Right-click and choose "Update Field" to generate the table of contents.',
    'running_header': 'NuFi — End-User Functional Specification',
    'footer': 'Confidential — NuFi QA',
}

if __name__ == '__main__':
    lang = sys.argv[1] if len(sys.argv) > 1 else 'en'
    base = Path(__file__).parent
    if lang == 'vi':
        from strings_vi import STRINGS_VI
        build(base / 'sections-vi', base / 'NuFi_Chat_Console_Spec_VI.docx', STRINGS_VI, MANIFEST)
    else:
        build(base / 'sections', base / 'NuFi_Chat_Console_Spec_EN.docx', STRINGS_EN, MANIFEST)
